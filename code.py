import os
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics.pairwise import cosine_similarity
import spacy
from sentence_transformers import SentenceTransformer
import numpy as np

# Load spaCy model
nlp = spacy.load('en_core_web_sm')

# Read documents from .docx files
def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Preprocess text (e.g., remove stop words, tokenize, lemmatize)
def preprocess_text(text):
    doc = nlp(text)
    return " ".join([token.lemma_ for token in doc if not token.is_stop and not token.is_punct])

# Convert documents to TF-IDF vectors
def vectorize_documents(docs):
    vectorizer = TfidfVectorizer(stop_words='english')
    return vectorizer.fit_transform(docs)

# Cluster documents
def cluster_documents(vectors, num_clusters=5):
    model = KMeans(n_clusters=num_clusters, random_state=42)
    model.fit(vectors)
    return model.predict(vectors)

# Remove redundant sentences or paragraphs (based on cosine similarity)
def remove_redundancy(cluster):
    sentences = cluster.split("\n")
    unique_sentences = []
    
    # Using sentence embeddings for similarity detection
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    embeddings = model.encode(sentences)
    
    for i, emb1 in enumerate(embeddings):
        if all(cosine_similarity([emb1], [emb2])[0][0] < 0.85 for emb2 in embeddings[:i]):
            unique_sentences.append(sentences[i])
    
    return "\n".join(unique_sentences)

# Save the output to a .docx file
def save_docx(content, output_path):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(output_path)

# Main process
def process_documents(input_folder, num_clusters=5):
    doc_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]
    docs = [preprocess_text(read_docx(os.path.join(input_folder, f))) for f in doc_files]
    vectors = vectorize_documents(docs)
    clusters = cluster_documents(vectors, num_clusters)

    # Organize documents by clusters and remove redundancy
    for i in range(num_clusters):
        cluster_docs = [docs[j] for j in range(len(docs)) if clusters[j] == i]
        consolidated_text = "\n".join([remove_redundancy(doc) for doc in cluster_docs])
        output_filename = f"Cluster_{i+1}_Consolidated_Output.docx"
        save_docx(consolidated_text, output_filename)

# Run the process
process_documents("path/to/input/folder")
